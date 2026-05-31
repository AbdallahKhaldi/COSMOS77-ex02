"""Programmatically fill the lecturer's cover template and convert to PDF.

Ported from HW1. Phase 12 retargets it for HW2 (exercise number = 2 and the
ex02 repo URL) and adds the matching `--exercise-number` flag + test. Until
then it is committed as-is so the tooling is in place. The lecturer's rule:
do NOT move template fields or add text outside them — python-docx field-fill
(append a run to the label's paragraph) satisfies this.
"""

from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from collections.abc import Iterable
from pathlib import Path

_FIELD_PLAN = (
    ("Submitting an exercise number", "1"),
    ("Group ID code", "COSMOS77"),
    ("Recommendation for self-scoring", None),  # placeholder, replaced at runtime
)


def fill_docx(template: Path, output_docx: Path, self_score: int) -> Path:
    """Insert the project's field values into `template` and save to `output_docx`."""
    from docx import Document

    doc = Document(str(template))
    paras = list(doc.paragraphs)
    student1 = ("212389712", "Abdallah", "Khaldi", "עבדאללה", "חאלדי")
    student2 = ("323118794", "Tasneem", "Natour", "תסנים", "נאטור")
    for para in paras:
        text = para.text.strip()
        if text.startswith("Submitting an exercise number"):
            _append_run(para, " 1")
        elif text.startswith("Group ID code"):
            _append_run(para, " COSMOS77")
        elif text.startswith("Recommendation for self-scoring"):
            _append_run(para, f" {self_score}")
        elif text.startswith("Link to GITHUB"):
            _append_run(para, " https://github.com/AbdallahKhaldi/COSMOS77-ex01")
        elif text.startswith("A late submission confirmation"):
            _append_run(para, " no")
    _fill_student(paras, header="Student 1", values=student1)
    _fill_student(paras, header="Student 2", values=student2)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


def _append_run(paragraph, text: str) -> None:  # noqa: ANN001
    run = paragraph.add_run(text)
    if paragraph.runs and paragraph.runs[0].font.size:
        run.font.size = paragraph.runs[0].font.size


def _fill_student(paragraphs: Iterable, *, header: str, values: tuple[str, ...]) -> None:  # noqa: ANN001
    paras = list(paragraphs)
    for idx, para in enumerate(paras):
        if para.text.strip() == header:
            labels = (
                "ID card",
                "First name in English",
                "Last name in English",
                "First name in Hebrew",
                "Last name in Hebrew",
            )
            for offset, (label, value) in enumerate(zip(labels, values, strict=True), start=1):
                target = paras[idx + offset]
                if target.text.strip().startswith(label):
                    _append_run(target, f" {value}")
            return


def convert_to_pdf(input_docx: Path, output_pdf: Path) -> Path:
    """Convert `input_docx` to PDF using docx2pdf, then LibreOffice fallback."""
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    if _try_docx2pdf(input_docx, output_pdf):
        return output_pdf
    if _try_libreoffice(input_docx, output_pdf):
        return output_pdf
    raise RuntimeError(
        "no PDF tool available. Install Microsoft Word (for docx2pdf) or "
        "`brew install --cask libreoffice` and re-run."
    )


def _try_docx2pdf(input_docx: Path, output_pdf: Path) -> bool:
    try:
        from docx2pdf import convert  # type: ignore[import-not-found]
    except ImportError:
        return False
    convert(str(input_docx), str(output_pdf))
    return output_pdf.exists()


def _try_libreoffice(input_docx: Path, output_pdf: Path) -> bool:
    for tool in ("soffice", "libreoffice"):
        if shutil.which(tool) is None:
            continue
        try:
            subprocess.run(
                [
                    tool,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_pdf.parent),
                    str(input_docx),
                ],
                check=True,
                capture_output=True,
            )
        except subprocess.CalledProcessError:
            continue
        produced = output_pdf.parent / (input_docx.stem + ".pdf")
        if produced.exists() and produced != output_pdf:
            produced.replace(output_pdf)
        return output_pdf.exists()
    return False


def verify_pdf(path: Path) -> None:
    """Validate the produced file is a real PDF (`%PDF` magic bytes)."""
    with path.open("rb") as fh:
        head = fh.read(5)
    if not head.startswith(b"%PDF-"):
        raise RuntimeError(f"{path} is not a valid PDF (got {head!r})")


def main(argv: list[str] | None = None) -> int:
    """Fill the cover template, convert to PDF, and verify the output."""
    parser = argparse.ArgumentParser(description="Generate the COSMOS77-ex02 cover PDF.")
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--self-score", type=int, default=85)
    args = parser.parse_args(argv)
    docx_out = args.output.with_suffix(".filled.docx")
    fill_docx(args.template, docx_out, self_score=args.self_score)
    convert_to_pdf(docx_out, args.output)
    verify_pdf(args.output)
    print(f"OK: cover PDF written to {args.output.resolve()}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    sys.exit(main())
